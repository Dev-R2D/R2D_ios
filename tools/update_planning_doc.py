from copy import deepcopy
from pathlib import Path
from docx import Document

SOURCE = Path("docs/R2D_개발팀용_통합서비스_상세기획서_v2.0.docx")
OUTPUT = Path("docs/R2D_개발팀용_통합서비스_상세기획서_v2.1.docx")


def set_text(paragraph, text):
    if paragraph.runs:
        paragraph.runs[0].text = text
        for run in paragraph.runs[1:]:
            run.text = ""
    else:
        paragraph.add_run(text)


def insert_after(paragraph, text, style=None):
    new_p = deepcopy(paragraph._p)
    for child in list(new_p):
        if child.tag.endswith("}r") or child.tag.endswith("}hyperlink"):
            new_p.remove(child)
    paragraph._p.addnext(new_p)
    from docx.text.paragraph import Paragraph
    result = Paragraph(new_p, paragraph._parent)
    if style is not None:
        result.style = style
    set_text(result, text)
    return result


doc = Document(SOURCE)

# Front matter and changelog.
doc.tables[0].cell(0, 1).text = "v2.1"
doc.tables[0].cell(1, 1).text = "2026.08.05"
row = doc.tables[2].add_row().cells
for cell, value in zip(row, [
    "v2.1",
    "2026.08.05",
    "KakaoMapsSDK UI Adapter 결정, 경로 검색 책임 분리, 네이티브 앱 키·무료 쿼터·fallback 운영 기준 보완",
    "모바일·백엔드·보안·QA",
]):
    cell.text = value

replacements = {
    "외부 지도·경로 공급자는 교체 가능하도록 Adapter 계층으로 감싼다.":
        "지도 렌더러와 경로 공급자는 서로 다른 Adapter로 분리한다. 지도 SDK는 UI Layer의 IMapRenderer로, 경로 검색은 IRouteRepository로 감싸며 NavigationEngine은 특정 지도 SDK를 참조하지 않는다.",
    "실제 외부 경로 API 연동과 안전 우선 경로 재정렬":
        "실제 외부 경로 API 연동과 안전 우선 경로 재정렬. KakaoMapsSDK 지도 렌더링과 경로 검색 API는 별도 계약으로 취급한다.",
    "지도 공급자 오류: 최소 경로 라인·텍스트 모드 제공":
        "지도 SDK 초기화·인증·쿼터 오류: MapKit fallback 또는 최소 경로 라인·텍스트 모드를 제공하고 진단 코드를 기록",
    "P0는 사전 정의 테스트 경로 또는 단일 외부 경로를 사용할 수 있음":
        "P0 데모는 사전 정의 테스트 경로를 사용한다. KakaoMapsSDK는 경로를 그리지만 경로 검색 결과를 생성하지 않으므로 실제 검색은 별도 IRouteRepository Adapter가 담당",
    "모바일 로컬 큐에는 인증 토큰과 정밀 위치가 포함되므로 플랫폼 보안 저장소와 암호화 DB를 사용한다.":
        "모바일 로컬 큐에는 인증 토큰과 정밀 위치가 포함되므로 플랫폼 보안 저장소와 암호화 DB를 사용한다. 지도 네이티브 앱 키는 소스·공유 Scheme·문서에 직접 기록하지 않고 Git 제외 xcconfig 또는 CI Secret으로 주입하며 REST/Admin 키를 모바일 앱에 포함하지 않는다.",
}

for paragraph in doc.paragraphs:
    if paragraph.text in replacements:
        set_text(paragraph, replacements[paragraph.text])

# Navigator home contract now names the renderer without conflating it with route search.
doc.tables[18].cell(4, 1).text = "GET /map/cells?bbox&zoom, GET /game/bosses/nearby, KakaoMapsSDK base map, cached risk layer"

# Close DEC-01 based on the actual P0 demo decision while preserving the future route-provider decision.
doc.tables[96].cell(1, 1).text = "P0 지도 렌더러/경로 공급 방식"
doc.tables[96].cell(1, 2).text = "결정: iOS KakaoMapsSDK Renderer + MapKit fallback / DemoRouteRepository fixture. 운영 경로 API 공급자는 별도 결정"
doc.tables[96].cell(1, 4).text = "P0 결정 완료, 운영 Provider 연동 전 재검토"

# Add the missing SDK responsibility and operational rules immediately after section 8.4.
anchor = next(p for p in doc.paragraphs if p.text == "8.4 지도 시각 규칙")
cursor = anchor
for text in [
    "8.5 지도 SDK Adapter 및 운영 기준",
    "iOS 기본 지도 표현은 KakaoMapsSDK v2를 사용한다. 단, SDK 또는 네이티브 앱 키가 준비되지 않은 Preview·Test·복구 환경에서는 Apple MapKit Renderer를 fallback으로 유지한다.",
    "KakaoMapRenderer는 R2DUI에서 Route, RiskCell, Turn, Destination, Current Location, Camera UI 모델만 변환한다. R2DCore와 NavigationEngine에는 KakaoMapsSDK 또는 MapKit import를 금지한다.",
    "KakaoMapsSDK는 지도 타일과 오버레이 렌더링을 담당하며 목적지 검색·경로 계산 API가 아니다. Route Search/Refresh/Cancel은 IRouteRepository가 담당하고, Route Matching·ETA·Off Route·ReRoute는 NavigationEngine이 담당한다.",
    "Kakao Developers 앱의 카카오맵 사용 상태와 무료 쿼터를 운영 체크리스트로 관리한다. 앱 초기화에는 네이티브 앱 키만 사용하고, Bundle ID app.r2d.mobile 일치 여부를 빌드·릴리스 전에 검증한다.",
    "키 미설정·placeholder·SDK 인증 실패·일간 쿼터 초과 시 앱이 종료되지 않아야 한다. 지도 fallback 또는 제한 모드로 전환하고 Route/Turn/ETA 텍스트 안내와 Safety Overlay를 유지한다.",
]:
    cursor = insert_after(cursor, text, anchor.style)

# Add a concrete current mobile baseline without rewriting the broader roadmap.
anchor = next(p for p in doc.paragraphs if p.text == "3.3 P2 또는 명시적 비범위")
cursor = anchor
for text in [
    "3.4 2026.08.05 모바일 구현 기준선",
    "현재 iOS 클라이언트는 Ride/Navigation Domain, Telemetry Queue·Upload, Progress Sync, Route/Risk Repository와 Cache, RoadWarningEngine, Authentication·Token Refresh, IMapRenderer 및 MapKit Adapter까지 구현되어 있다.",
    "이번 v2.1에서 KakaoMapsSDK는 기존 IMapRenderer의 Production Adapter로 추가한다. Demo/Preview/Test의 Mock·Fake 구조와 기존 Domain 계약은 변경하지 않는다.",
    "Navigator 데모는 fixture Route/Risk/Progress와 Replay를 사용하므로 서버·실제 GPS 없이 재현 가능하다. 운영 Route Provider, Road Cell Server, Game Server 연결 상태와 데모 완료 상태를 혼동하지 않는다.",
]:
    cursor = insert_after(cursor, text, anchor.style)

doc.core_properties.title = "R2D 개발팀용 통합 서비스 상세 기획서 v2.1"
doc.core_properties.comments = "v2.1: KakaoMapsSDK Adapter 및 지도/경로 책임 분리 보완"
doc.save(OUTPUT)
print(OUTPUT)
